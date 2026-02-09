"""Document processor for chunking and parsing files.

Enhanced with universal document metadata extraction for various document types:
- Academic papers, Corporate documents, PPT, Meeting minutes,
- Legal documents, Insurance policies, and more.
"""

import logging
import os
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

import aiofiles
from langchain_core.documents import Document as LangchainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import settings
from app.services.rag.document_metadata import (
    detect_document_type,
    extract_universal_metadata,
    DocumentType,
)

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Process documents for RAG indexing.

    Supports: PDF, TXT, Markdown, DOCX, CSV
    Performs: Loading, text extraction, chunking
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".csv", ".json"}

    def __init__(
        self,
        chunk_size: int = settings.chunk_size,
        chunk_overlap: int = settings.chunk_overlap,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            length_function=len,
            is_separator_regex=False,
        )

    async def process_file(
        self,
        file_path: str,
        file_name: str,
        user_id: str,
        document_id: str,
        project_id: Optional[str] = None,
    ) -> list[LangchainDocument]:
        """
        Process a file and return chunked documents.

        Args:
            file_path: Path to the file
            file_name: Original file name
            user_id: Owner user ID
            document_id: Document ID in database
            project_id: Optional project ID

        Returns:
            List of LangChain documents with content and metadata
        """
        ext = Path(file_path).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        # Load document based on type
        documents = await self._load_document(file_path, ext, file_name)

        if not documents:
            raise ValueError(f"Failed to load documents from {file_name}")

        # Check for empty content
        if all(not doc.page_content.strip() for doc in documents):
            raise ValueError(f"Document appears to be empty: {file_name}")

        # Split into chunks
        chunks = self.text_splitter.split_documents(documents)

        # Add metadata
        file_metadata = {
            "file_name": file_name,
            "file_type": ext[1:],  # Remove dot
            "document_id": document_id,
            "user_id": user_id,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        if project_id:
            file_metadata["project_id"] = project_id

        for i, chunk in enumerate(chunks):
            chunk.metadata.update(file_metadata)
            chunk.metadata["chunk_index"] = i
            chunk.metadata["page"] = chunk.metadata.get("page", 0) + 1

        logger.info(f"Processed {file_name}: {len(chunks)} chunks created")
        return chunks

    async def process_file_content(
        self,
        content: bytes,
        file_name: str,
        user_id: str,
        document_id: str,
        project_id: Optional[str] = None,
    ) -> list[LangchainDocument]:
        """
        Process file content from bytes.

        Args:
            content: File content as bytes
            file_name: Original file name
            user_id: Owner user ID
            document_id: Document ID
            project_id: Optional project ID

        Returns:
            List of chunked LangChain documents
        """
        ext = Path(file_name).suffix.lower()

        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
            temp_file.write(content)
            temp_path = temp_file.name

        try:
            return await self.process_file(
                file_path=temp_path,
                file_name=file_name,
                user_id=user_id,
                document_id=document_id,
                project_id=project_id,
            )
        finally:
            # Clean up temp file
            os.unlink(temp_path)

    async def _load_document(
        self,
        file_path: str,
        ext: str,
        file_name: str = "",
    ) -> list[LangchainDocument]:
        """Load document based on file type."""
        try:
            if ext == ".pdf":
                return await self._load_pdf(file_path, file_name)
            elif ext in {".txt", ".md"}:
                return await self._load_text(file_path)
            elif ext == ".docx":
                return await self._load_docx(file_path)
            elif ext == ".csv":
                return await self._load_csv(file_path)
            elif ext == ".json":
                return await self._load_json(file_path)
            else:
                raise ValueError(f"Unsupported extension: {ext}")
        except Exception as e:
            logger.error(f"Failed to load {file_path}: {e}")
            raise

    async def _load_pdf(
        self,
        file_path: str,
        file_name: str = "",
        extract_metadata: bool = True,
    ) -> list[LangchainDocument]:
        """
        Load PDF document using PyMuPDF with universal metadata extraction.

        Args:
            file_path: Path to PDF file
            file_name: Original file name (for document type detection)
            extract_metadata: Whether to extract metadata from first page

        Returns:
            List of LangChain documents with content and metadata
        """
        import fitz  # PyMuPDF

        documents = []
        doc_metadata = {}
        doc = fitz.open(file_path)

        # Collect first few pages for metadata extraction
        first_pages_text = ""
        for page_num in range(min(3, len(doc))):  # First 3 pages
            page = doc[page_num]
            first_pages_text += page.get_text() + "\n"

        for page_num, page in enumerate(doc):
            text = page.get_text()
            if not text.strip():
                continue

            # Extract universal metadata from first pages
            if page_num == 0 and extract_metadata:
                doc_metadata = extract_universal_metadata(first_pages_text, file_name)

                # Get detected document type
                doc_type = doc_metadata.get("doc_type", "general")
                logger.info(f"Detected document type: {doc_type} for {file_name}")

                # If abstract exists, prepend to first page for better context
                abstract = doc_metadata.get("abstract")
                if abstract:
                    text = f"[Abstract]\n{abstract}\n\n{text}"

            documents.append(
                LangchainDocument(
                    page_content=text,
                    metadata={
                        "page": page_num,
                        "source": file_path,
                        **doc_metadata,
                    },
                )
            )

        doc.close()
        return documents

    async def process_pdf_with_images(
        self,
        file_path: str,
        file_name: str,
        user_id: str,
        document_id: str,
        project_id: Optional[str] = None,
    ) -> dict[str, Any]:
        """
        Process PDF including image extraction.

        Args:
            file_path: Path to the PDF file
            file_name: Original file name
            user_id: Owner user ID
            document_id: Document ID
            project_id: Optional project ID

        Returns:
            Processing statistics including image count
        """
        try:
            # Process text chunks
            chunks = await self.process_file(
                file_path=file_path,
                file_name=file_name,
                user_id=user_id,
                document_id=document_id,
                project_id=project_id,
            )

            stats = self.get_chunk_stats(chunks)

            # Process images
            try:
                from app.services.rag.image.processor import ImageProcessor

                image_processor = ImageProcessor()
                image_stats = await image_processor.process_pdf_images(
                    pdf_path=file_path,
                    user_email=user_id,
                    file_name=file_name,
                    project_id=project_id or "",
                )
                stats["images_found"] = image_stats.get("images_found", 0)
                stats["images_indexed"] = image_stats.get("images_indexed", 0)
            except ImportError:
                logger.warning("Image processor dependencies not available")
                stats["images_found"] = 0
                stats["images_indexed"] = 0
            except Exception as e:
                logger.warning(f"Image processing failed: {e}")
                stats["images_found"] = 0
                stats["images_indexed"] = 0

            return {
                "chunks": chunks,
                "stats": stats,
            }

        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise

    async def _load_text(self, file_path: str) -> list[LangchainDocument]:
        """Load text or markdown file."""
        async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
            content = await f.read()

        return [
            LangchainDocument(
                page_content=content,
                metadata={"source": file_path},
            )
        ]

    async def _load_docx(self, file_path: str) -> list[LangchainDocument]:
        """Load DOCX document."""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        content = "\n\n".join(full_text)

        return [
            LangchainDocument(
                page_content=content,
                metadata={"source": file_path},
            )
        ]

    async def _load_csv(self, file_path: str) -> list[LangchainDocument]:
        """Load CSV file."""
        import csv

        documents = []

        async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
            content = await f.read()

        reader = csv.DictReader(content.splitlines())
        for i, row in enumerate(reader):
            # Convert row to text
            text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
            if text.strip():
                documents.append(
                    LangchainDocument(
                        page_content=text,
                        metadata={"source": file_path, "row": i},
                    )
                )

        return documents

    async def _load_json(self, file_path: str) -> list[LangchainDocument]:
        """Load JSON file."""
        import json

        async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
            content = await f.read()

        data = json.loads(content)

        # Convert to text representation
        if isinstance(data, list):
            documents = []
            for i, item in enumerate(data):
                text = json.dumps(item, ensure_ascii=False, indent=2)
                documents.append(
                    LangchainDocument(
                        page_content=text,
                        metadata={"source": file_path, "index": i},
                    )
                )
            return documents
        else:
            text = json.dumps(data, ensure_ascii=False, indent=2)
            return [
                LangchainDocument(
                    page_content=text,
                    metadata={"source": file_path},
                )
            ]

    def get_chunk_stats(self, documents: list[LangchainDocument]) -> dict[str, Any]:
        """Get statistics about chunked documents."""
        if not documents:
            return {"chunk_count": 0, "total_chars": 0, "avg_chunk_size": 0}

        total_chars = sum(len(doc.page_content) for doc in documents)
        return {
            "chunk_count": len(documents),
            "total_chars": total_chars,
            "avg_chunk_size": total_chars // len(documents),
            "min_chunk_size": min(len(doc.page_content) for doc in documents),
            "max_chunk_size": max(len(doc.page_content) for doc in documents),
        }
